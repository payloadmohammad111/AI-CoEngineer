"""
Module: exporter
Provides functions to export generated documents (SRS, design docs, test plans, etc.)
with proper formatting for PDF, DOCX, Markdown, HTML, and JSON.
"""

import json
from typing import Dict, Any

# ======================
# PDF Export (ReportLab)
# ======================
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

def export_to_pdf(content: Dict[str, Any], filename: str) -> None:
    """
    Export structured content to a nicely formatted PDF file.

    Args:
        content (dict): Structured content (sections, lists, etc.).
        filename (str): The name of the output PDF file.
    """
    try:
        doc = SimpleDocTemplate(filename, pagesize=letter)
        styles = getSampleStyleSheet()
        flowables = []

        for section, value in content.items():
            flowables.append(Paragraph(f"<b>{section}</b>", styles['Heading2']))
            flowables.append(Spacer(1, 12))

            if isinstance(value, list):
                for item in value:
                    flowables.append(Paragraph(f"- {item}", styles['Normal']))
            elif isinstance(value, dict):
                for subkey, subval in value.items():
                    flowables.append(Paragraph(f"<b>{subkey}:</b> {subval}", styles['Normal']))
            else:
                flowables.append(Paragraph(str(value), styles['Normal']))

            flowables.append(Spacer(1, 12))

        doc.build(flowables)
        print(f"✅ PDF file '{filename}' created successfully.")
    except Exception as e:
        print(f"❌ Error exporting to PDF: {e}")


# ======================
# DOCX Export (python-docx)
# ======================
from docx import Document

def export_to_docx(content: Dict[str, Any], filename: str) -> None:
    """
    Export structured content to a nicely formatted DOCX file.

    Args:
        content (dict): Structured content.
        filename (str): The name of the output DOCX file.
    """
    try:
        doc = Document()
        doc.add_heading("Software Requirements Specification (SRS)", 0)

        for section, value in content.items():
            doc.add_heading(section, level=1)

            if isinstance(value, list):
                for item in value:
                    doc.add_paragraph(item, style="ListBullet")
            elif isinstance(value, dict):
                for subkey, subval in value.items():
                    doc.add_paragraph(f"{subkey}: {subval}", style="Normal")
            else:
                doc.add_paragraph(str(value), style="Normal")

        doc.save(filename)
        print(f"✅ DOCX file '{filename}' created successfully.")
    except Exception as e:
        print(f"❌ Error exporting to DOCX: {e}")


# ======================
# Markdown Export
# ======================
def export_to_markdown(content: Dict[str, Any], filename: str) -> None:
    """
    Export structured content to a Markdown file.
    """
    try:
        with open(filename, "w", encoding="utf-8") as md:
            md.write("# Software Requirements Specification (SRS)\n\n")
            for section, value in content.items():
                md.write(f"## {section}\n")
                if isinstance(value, list):
                    for item in value:
                        md.write(f"- {item}\n")
                elif isinstance(value, dict):
                    for subkey, subval in value.items():
                        md.write(f"- **{subkey}:** {subval}\n")
                else:
                    md.write(f"{value}\n")
                md.write("\n")
        print(f"✅ Markdown file '{filename}' created successfully.")
    except Exception as e:
        print(f"❌ Error exporting to Markdown: {e}")


# ======================
# HTML Export
# ======================
def export_to_html(content: Dict[str, Any], filename: str) -> None:
    """
    Export structured content to an HTML file.
    """
    try:
        with open(filename, "w", encoding="utf-8") as html:
            html.write("<html><head><title>SRS</title></head><body>")
            html.write("<h1>Software Requirements Specification (SRS)</h1>")
            for section, value in content.items():
                html.write(f"<h2>{section}</h2>")
                if isinstance(value, list):
                    html.write("<ul>")
                    for item in value:
                        html.write(f"<li>{item}</li>")
                    html.write("</ul>")
                elif isinstance(value, dict):
                    html.write("<ul>")
                    for subkey, subval in value.items():
                        html.write(f"<li><b>{subkey}:</b> {subval}</li>")
                    html.write("</ul>")
                else:
                    html.write(f"<p>{value}</p>")
            html.write("</body></html>")
        print(f"✅ HTML file '{filename}' created successfully.")
    except Exception as e:
        print(f"❌ Error exporting to HTML: {e}")


# ======================
# JSON Export
# ======================
def export_to_json(content: Dict[str, Any], filename: str) -> None:
    """
    Export structured content to a JSON file.
    """
    try:
        with open(filename, "w", encoding="utf-8") as js:
            json.dump(content, js, indent=4, ensure_ascii=False)
        print(f"✅ JSON file '{filename}' created successfully.")
    except Exception as e:
        print(f"❌ Error exporting to JSON: {e}")
